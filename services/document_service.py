import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import current_app

class DocumentService:
    """Сервис для генерации документов"""
    
    def generate_school_request(self, application_id):
        """
        Генерация письма запроса в школу
        
        Args:
            application_id (int): ID заявки
            
        Returns:
            str: Путь к сгенерированному документу
        """
        from models.application import Application
        
        application = Application.query.get(application_id)
        if not application:
            return None
            
        # Создаем документ
        doc = Document()
        
        # Настраиваем поля страницы
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1.5)
        
        # Добавляем заголовок
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.add_run('ЗАПРОС ИНФОРМАЦИИ')
        header_run.bold = True
        header_run.font.size = Pt(16)
        
        # Добавляем дату
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_paragraph.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
        
        # Добавляем обращение
        doc.add_paragraph().add_run(f"Уважаемый руководитель {application.school.name}!")
        
        # Добавляем текст запроса
        text = doc.add_paragraph()
        text.add_run(
            f"В рамках проекта по сбору информации о выпускниках и учителях, "
            f"выпускник Вашей школы {application.graduate.full_name} "
            f"указал, что обучался в Вашей школе с {application.start_year} по {application.end_year} год.\n\n"
            f"Просим Вас подтвердить данную информацию и предоставить список учителей, "
            f"которые работали в школе в указанный период.\n\n"
            f"Для заполнения информации, пожалуйста, перейдите по следующей ссылке:"
        )
        
        # Добавляем ссылку
        from services.link_service import LinkService
        link_service = LinkService()
        url = link_service.generate_school_link(application.id)
        
        link_paragraph = doc.add_paragraph()
        link_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        link_run = link_paragraph.add_run(url)
        link_run.font.size = Pt(12)
        link_run.bold = True
        
        # Добавляем QR-код
        from services.qr_service import QRService
        qr_service = QRService()
        qr_path = qr_service.generate_school_qr(application.id)
        
        if qr_path:
            doc.add_paragraph().add_run("Или отсканируйте QR-код:")
            doc.add_picture(qr_path, width=Cm(5))
        
        # Добавляем подпись
        doc.add_paragraph().add_run("\n\nС уважением,\nАдминистрация проекта")
        
        # Сохраняем документ
        filename = f"request_school_{application_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
        doc.save(file_path)
        
        return file_path
    
    def save_document(self, document, filename):
        """
        Сохранение документа на сервере
        
        Args:
            document: Документ для сохранения
            filename (str): Имя файла
            
        Returns:
            str: Путь к сохраненному документу
        """
        file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
        document.save(file_path)
        return file_path
    
    def get_document_path(self, filename):
        """
        Получение пути к документу
        
        Args:
            filename (str): Имя файла
            
        Returns:
            str: Полный путь к документу
        """
        return os.path.join(current_app.config['UPLOAD_FOLDER'], filename)